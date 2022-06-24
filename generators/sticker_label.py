import textwrap
import logging
from io import BytesIO
from pathlib import Path
from config import DEFAULT_FONT
from utilities import Utils as utils

from PIL import Image
from PIL import ImageFont, ImageDraw

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class StickerLabelGenerator():
    def __init__(self):
        self.label_width, self.label_height = (342, 1080)
        self.barcode_width = self.label_width
        self.font_size = 16
        self.STICKERS = {}
        self.STICKERS_ALONE = {}
        self.SAVE_DIR = f"userdata"


    def clear(self, uid: str|int, temp_files: bool = False):
        if temp_files:
            for file in Path(self.SAVE_DIR, str(uid)).glob("stickers_*"):
                file.unlink()
        else:
            del self.STICKERS[uid]
            try:
                del self.STICKERS_ALONE[uid]
            except:
                logging.info(f"Can't delete STICKERS_ALONE (id:{uid})")


    def label_create(self, uid: str|int, product_params:dict, barcode_img:BytesIO) -> Path:
        FILENAME = f"label_{product_params['Артикул']}.png"

        barcode_png = Image.open(barcode_img)
        w, h = barcode_png.size

        barcode_height = int(self.barcode_width * h / w)
        barcode_png = barcode_png.resize((
            self.barcode_width, barcode_height
        ), Image.ANTIALIAS)

        img = Image.new(
            'RGB',
            (self.label_width, self.label_height),
            (255, 255, 255)  # White
        )
        img.paste(
            barcode_png, (
                (self.label_width - self.barcode_width) // 2,
                0
            )
        )
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype(DEFAULT_FONT, size=self.font_size)

        i = barcode_height
        for k, v in product_params.items():
            if k not in ['Штрихкод']:
                textwrapped = textwrap.wrap(f"{k}: {v}", width=31)
                offset = (35, i)
                text = '\n'.join(textwrapped)
                draw.text((offset), text, font=font, fill="black")
                newlines = text.count("\n")+1
                i = i+(20*(newlines if newlines > 0 else 1)) # отступы ебашим

        savepath = Path(self.SAVE_DIR, str(uid), FILENAME)
        img = img.crop((0,0, self.label_width, i+5))
        img.save(savepath, quality=100)
        logging.info(f"USER'S (id: {uid}) NEW LABEL CREATED")

        return savepath


    def sticker_add(self, uid: str|int, user_code_png: BytesIO, should_double:bool = False) -> None:
        user_CodePng = Image.open(user_code_png)
        try:
            self.STICKERS[uid].append([user_CodePng])
        except KeyError:
            self.STICKERS[uid] = [[user_CodePng]]

        if should_double:
            try:
                self.STICKERS_ALONE[uid] += 1
            except KeyError:
                self.STICKERS_ALONE[uid] = 1


    def complete_file(self, uid: str|int, articul:str) -> Path:
        logging.info(f"USER (id: {uid}) GENERATING STICKERS")
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        rows = len(self.STICKERS[uid])
        descriptions = False
        
        if uid in self.STICKERS_ALONE:
            rows += self.STICKERS_ALONE[uid] // 2
            descriptions = self.STICKERS_ALONE[uid]
    
        if rows == 1:   # а хули оно ломается
            rows = 2
        table = document.add_table(rows=rows, cols=3)
    
        barcode_png = Image.open(Path(self.SAVE_DIR, str(uid), f"label_{articul}.png"))

        # ниже пиздец который отнял пол жизни
        # за такой код убивают.............................
        # создано в содружестве с @moksempython
        bills = self.STICKERS[uid] 
        row, column = 0, 0
        while len(bills) > 0:
            for enum1, elem in enumerate(bills):
                if column == 3:# (ну типа по три колонки или хз):
                    column = 0
                    row += 2
                for image in elem:
                    # barcode -- описание
                    row_cells = table.rows[row].cells
                    barcode_cell = row_cells[column].paragraphs[0]
                    barcode_pr = barcode_cell.add_run()
                    # usercode -- штрихкод отправки
                    row_cells = table.rows[row+1].cells
                    usercode_cell = row_cells[column].paragraphs[0]
                    usercode_pr = usercode_cell.add_run()
                    if image.size[0] == 236:
                        usercode_pr.add_picture(utils.image2file_bytes(image), width=Cm(4), height=Cm(3))
                        barcode_pr.add_picture(utils.image2file_bytes(barcode_png), width=Cm(4.82))
                        barcode_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif image.size[0] == 342 or image.size[0] == 320:
                        usercode_pr.add_picture(utils.image2file_bytes(image), width=Cm(5.8), height=Cm(4))
                        barcode_pr.add_picture(utils.image2file_bytes(barcode_png), width=Cm(5.82))
                        barcode_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        raise Exception("Неподдерживаемый формат штрихкода!")
                    usercode_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                del bills[enum1]
                column += 1
        # А ТУТ добавляем маркировки товаров (описание) без клиентских штрихов если попросили...
        state1 = True # это единственный адекватный вариант который я придумал
        if descriptions:
            while descriptions > 0:
                if column == 3:# (ну типа по три колонки или хз):
                    column = 0
                    row += 2
                if state1:
                    row_cells = table.rows[row].cells
                    barcode_cell_first = row_cells[column].paragraphs[0]
                    barcode_pr_first = barcode_cell_first.add_run()
                    
                    barcode_pr_first.add_picture(utils.image2file_bytes(barcode_png), width=Cm(5.82))
                    state1 = False
                    barcode_cell_first.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells = table.rows[row+1].cells
                    barcode_cell_second = row_cells[column].paragraphs[0]
                    barcode_pr_second = barcode_cell_second.add_run()
                    
                    barcode_pr_second.add_picture(utils.image2file_bytes(barcode_png), width=Cm(5.82))
                    barcode_cell_second.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    state1 = True
                    column += 1

                descriptions -= 1
        filename = f"stickers_{utils.uuid_gen()}.docx"
        filepath = Path(self.SAVE_DIR, str(uid), filename)
        
        document.save(filepath)
        logging.info(f"USER (id: {uid}) GENERATED STICKERS SUCCESSFULLY")
        try:
            return filepath
        finally:
            self.clear(uid)
